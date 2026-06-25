#!/usr/bin/env python3
"""
HI2 IRI 解码报告生成器模板

用法:
  python3 gen_hi2_report.py decoded_a.html decoded_b.html -o report.docx

依赖:
  pip install python-docx

从 ETSI-ASN1-Assistant 生成的 HTML 解码文件中提取 HI2 IRI 数据，
生成结构化 Word 报告，含 PANI 位置分析、公网IP提取、来源标注。
"""

import re, json, sys, os
from collections import defaultdict
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ──────────────────────────────────────────────
# 辅助函数 (参见 SKILL.md PANI 提取部分)
# ──────────────────────────────────────────────

def set_cn_font(run, cn_font='微软雅黑', en_font='Arial', size=None, bold=None, color=None):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    if size: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if color: run.font.color.rgb = color

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = h
        set_cell_shading(cell, '1F4E79')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True; r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(255, 255, 255); set_cn_font(r)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            if ri % 2 == 1: set_cell_shading(cell, 'E8F0FE')
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9); set_cn_font(r)
    if col_widths:
        for cwi, w in enumerate(col_widths):
            for row in table.rows: row.cells[cwi].width = Cm(w)
    return table

def parse_pani(sip_msg):
    """从SIP消息中提取P-Access-Network-Info"""
    panis = []
    for m in re.finditer(r'P-Access-Network-Info:\s*([^\r\n]+)', sip_msg):
        raw = m.group(1).strip()
        info = {'raw': raw}
        at = re.search(r'^(\S+)', raw)
        if at: info['access_type'] = at.group(1)
        ci = re.search(r'utran-cell-id-3gpp=([0-9A-Fa-f]+)', raw)
        if ci: info['cell_id'] = ci.group(1)
        ue_ip = re.search(r'ue-ip=([0-9.]+)', raw)
        if ue_ip: info['ue_ip'] = ue_ip.group(1)
        ue_port = re.search(r'ue-port=(\d+)', raw)
        if ue_port: info['ue_port'] = ue_port.group(1)
        sbc = re.search(r'sbc-domain=([^;\"]+)', raw.replace('\"', ''))
        if sbc: info['sbc_domain'] = sbc.group(1).strip()
        # VoWiFi公网IP (注意区别于ue-ip)
        wlan_ip = re.search(r'Wlan-ue-local-ip=([0-9.]+)', raw)
        if wlan_ip: info['wlan_public_ip'] = wlan_ip.group(1)
        panis.append(info)
    return panis

# ──────────────────────────────────────────────
# 数据提取
# ──────────────────────────────────────────────

def extract_packets(html_path):
    with open(html_path, 'r') as f:
        content = f.read()
    blocks = re.findall(r'<pre><code class=\"json\">(\[[\s\S]*?\])</code></pre>', content)
    packets = []
    for block in blocks:
        try:
            data = json.loads(block)
            if data and isinstance(data, list) and len(data) > 0:
                record = data[0].get('iRI-Report-record', {})
                ims = record.get('imsGenIRIReport', {})
                sip = ims.get('sipMessage', '')
                packets.append({
                    'liid': record.get('lawfulInterceptionIdentifier', '') or '',
                    'timestamp': record.get('timeStamp', {}).get('localTime', {}).get('generalizedTime', ''),
                    'event': record.get('umts-Cs-Event', ''),
                    'direction': record.get('intercepted-Call-Direct', ''),
                    'cin': str(record.get('communicationIdentifier', {}).get('network-Identifier', {}).get('network-Element-Identifier', {}).get('e164-Format', '')),
                    'operator': record.get('communicationIdentifier', {}).get('network-Identifier', {}).get('operator-Identifier', ''),
                    'ims_charging_id': record.get('imsChargingID', ''),
                    'sip_msg': sip,
                    'pani': parse_pani(sip),
                    'cseq': re.search(r'CSeq:\s*\d+\s+(\S+)', sip).group(1) if re.search(r'CSeq:\s*\d+\s+(\S+)', sip) else '',
                    'user_agent': re.search(r'User-Agent:\s*([^\r\n]+)', sip).group(1).strip() if re.search(r'User-Agent:\s*([^\r\n]+)', sip) else '',
                })
        except:
            pass
    return packets

# ──────────────────────────────────────────────
# 主流程 (参考 SKILL.md 报告结构)
# ──────────────────────────────────────────────

def generate_report(html_files, output_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10.5)

    h = doc.add_heading('HI2 IRI 解码分析报告', level=1)
    for r in h.runs: set_cn_font(r, size=22, bold=True)
    doc.add_paragraph(f'来源文件: {", ".join(html_files)}')

    all_packets = []
    for f in html_files:
        all_packets.extend(extract_packets(f))

    sessions = defaultdict(list)
    for p in all_packets:
        if p['liid']: sessions[p['liid']].append(p)

    add_styled_table(doc,
        ['统计项', '数值'],
        [['解码报告总数', str(len(all_packets))],
         ['独立会话数', str(len(sessions))]],
        col_widths=[5, 8])

    doc.add_paragraph('\n注意: 每条结论须标注来源(PCAP文件名+报文序号)')
    doc.save(output_path)
    print(f'[✓] 报告已生成: {output_path}')

if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='生成HI2 IRI解码报告')
    parser.add_argument('input_files', nargs='+', help='ETSI-ASN1-Assistant输出的HTML文件')
    parser.add_argument('-o', '--output', default='HI2_Report.docx', help='输出DOCX路径')
    args = parser.parse_args()
    generate_report(args.input_files, args.output)
