#!/usr/bin/env python3
"""
模板：生成格式化周报 DOCX
=======================
使用方式：
  1. 修改下方 DATA 字典的周期、姓名、各章节内容
  2. python3 gen_weekly_docx.py
  3. 检查输出文件

依赖: python-docx (pip install python-docx)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ═══════════════════════ 数据区 ═══════════════════════
DATA = {
    "period": "2026-06-08 ~ 2026-06-12",
    "author": "毛恒振",
    "output_dir": os.path.expanduser("~/sinovatio/周报/"),
    "filename": "maohengzhen_周报_PERIOD.docx",  # PERIOD 会被替换

    # ── 一、核心完成 ──
    "core_items": [
        "事项1 — 效果/数据支撑",
        "事项2 — 效果/数据支撑",
        "事项3 — 效果/数据支撑",
    ],

    # ── 二、工作详情 ──
    # 每组: (标题, 占比, [(日期, 工作项, 进度, 说明), ...])
    "work_groups": [
        {
            "title": "组A名称",
            "ratio": "50%",
            "rows": [
                ["6/11", "工作项1", "100%", "说明"],
            ],
        },
    ],

    # ── 其他工作 ──
    "other_work": [
        ["日期", "工作项", "说明"],
    ],

    # ── 三、问题 ──
    "issues": [
        ["问题", "根因", "解决方案", "结果"],
    ],

    # ── 四、下周计划 ──
    "next_plan": [
        ["P0", "目标", "完成标准"],
    ],

    # ── 五、数据汇总 ──
    "summary": [
        ["指标", "值"],
    ],
}

# ═══════════════════════ 工具函数 ═══════════════════════
import os

def set_cell_shading(cell, color_hex):
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading_elm.append(shading)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    return h

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, '1A56DB')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for ri in range(len(rows) + 1):
            for ci, w in enumerate(col_widths):
                table.rows[ri].cells[ci].width = Cm(w)
    doc.add_paragraph()

# ═══════════════════════ 主函数 ═══════════════════════
def generate(data):
    doc = Document()

    # 字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题
    title = doc.add_heading(f'【周报】{data["period"]}', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{data["author"]} | Sinovatio 一营四区')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    # 一、核心完成
    add_heading(doc, '一、本周核心完成', level=1)
    for i, item in enumerate(data["core_items"], 1):
        p = doc.add_paragraph(f'{i}. {item}')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    # 二、工作详情
    add_heading(doc, '二、本周工作详情', level=1)
    for group in data["work_groups"]:
        add_heading(doc, f'【{group["title"]}】（占比 {group["ratio"]}）', level=2)
        if "headers" in group:
            add_table(doc, group["headers"], group["rows"], group.get("col_widths"))
        else:
            add_table(doc, group["rows"][0], group["rows"][1:], group.get("col_widths"))

    if data.get("other_work"):
        add_heading(doc, '【其他工作】', level=2)
        add_table(doc, data["other_work"][0], data["other_work"][1:], [2, 5, 8.5])

    # 三、问题
    add_heading(doc, '三、问题与解决', level=1)
    add_table(doc, data["issues"][0], data["issues"][1:], [4.5, 4, 4, 1.5])

    # 四、计划
    add_heading(doc, '四、下周计划', level=1)
    add_table(doc, data["next_plan"][0], data["next_plan"][1:], [2, 7, 6])

    # 五、数据
    add_heading(doc, '五、数据汇总', level=1)
    add_table(doc, data["summary"][0], data["summary"][1:], [5.5, 10])

    # 输出
    fn = data["filename"].replace("PERIOD", data["period"].replace(" ~ ", "-"))
    out_path = os.path.join(data["output_dir"], fn)
    os.makedirs(data["output_dir"], exist_ok=True)
    doc.save(out_path)
    print(f'✅ 周报已生成: {out_path}')

if __name__ == "__main__":
    generate(DATA)
